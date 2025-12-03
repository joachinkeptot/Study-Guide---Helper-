"""
Document Parser Service - Extract and structure content from various file formats.

This module handles file uploads, text extraction from multiple formats,
and LLM-based content structuring for study materials.
"""

import os
import io
import mimetypes
from typing import Dict, Any, Optional, BinaryIO, Union
from pathlib import Path
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParserError(Exception):
    """Base exception for document parsing errors."""
    pass


class UnsupportedFileTypeError(DocumentParserError):
    """Raised when file type is not supported."""
    pass


class TextExtractionError(DocumentParserError):
    """Raised when text extraction fails."""
    pass


class DocumentParser:
    """
    Service for parsing and extracting text from various document formats.
    
    Supported formats:
    - PDF (.pdf)
    - Word documents (.docx, .doc)
    - Text files (.txt, .md)
    - Images (.png, .jpg, .jpeg, .tiff, .bmp) - with OCR
    """
    
    # Supported file extensions and their MIME types
    SUPPORTED_EXTENSIONS = {
        '.pdf': ['application/pdf'],
        '.docx': ['application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
        '.doc': ['application/msword'],
        '.txt': ['text/plain'],
        '.md': ['text/markdown', 'text/plain'],
        '.png': ['image/png'],
        '.jpg': ['image/jpeg'],
        '.jpeg': ['image/jpeg'],
        '.tiff': ['image/tiff'],
        '.bmp': ['image/bmp']
    }
    
    def __init__(self, llm_provider=None):
        """
        Initialize the document parser.
        
        Args:
            llm_provider: LLM provider instance for content structuring.
                         If None, will use default provider when needed.
        """
        self.llm_provider = llm_provider
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required dependencies are available."""
        self.has_pdf_support = False
        self.has_docx_support = False
        self.has_ocr_support = False
        
        try:
            import pdfplumber
            self.has_pdf_support = True
            logger.info("PDF support enabled (pdfplumber)")
        except ImportError:
            try:
                import PyPDF2
                self.has_pdf_support = True
                logger.info("PDF support enabled (PyPDF2)")
            except ImportError:
                logger.warning("No PDF library found. Install pdfplumber or PyPDF2 for PDF support.")
        
        try:
            import docx
            self.has_docx_support = True
            logger.info("DOCX support enabled")
        except ImportError:
            logger.warning("python-docx not found. Install for Word document support.")
        
        try:
            import pytesseract
            from PIL import Image
            self.has_ocr_support = True
            logger.info("OCR support enabled")
        except ImportError:
            logger.warning("pytesseract or Pillow not found. Install for image OCR support.")
    
    def is_supported_file(self, filename: str, mime_type: Optional[str] = None) -> bool:
        """
        Check if a file type is supported.
        
        Args:
            filename: Name of the file
            mime_type: Optional MIME type of the file
            
        Returns:
            True if file type is supported, False otherwise
        """
        ext = Path(filename).suffix.lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            return False
        
        # Check if we have the required library for this file type
        if ext == '.pdf' and not self.has_pdf_support:
            return False
        if ext in ['.docx', '.doc'] and not self.has_docx_support:
            return False
        if ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp'] and not self.has_ocr_support:
            return False
        
        return True
    
    def extract_text(
        self,
        file_content: Union[bytes, BinaryIO],
        filename: str,
        mime_type: Optional[str] = None
    ) -> str:
        """
        Extract text content from a file.
        
        Args:
            file_content: File content as bytes or file-like object
            filename: Name of the file (used to determine type)
            mime_type: Optional MIME type hint
            
        Returns:
            Extracted text content
            
        Raises:
            UnsupportedFileTypeError: If file type is not supported
            TextExtractionError: If text extraction fails
        """
        if not self.is_supported_file(filename, mime_type):
            ext = Path(filename).suffix.lower()
            raise UnsupportedFileTypeError(
                f"Unsupported file type: {ext}. "
                f"Supported types: {', '.join(self.SUPPORTED_EXTENSIONS.keys())}"
            )
        
        ext = Path(filename).suffix.lower()
        
        # Convert to bytes if needed
        if hasattr(file_content, 'read'):
            file_bytes = file_content.read()
        else:
            file_bytes = file_content
        
        try:
            if ext == '.pdf':
                return self._extract_from_pdf(file_bytes)
            elif ext in ['.docx', '.doc']:
                return self._extract_from_docx(file_bytes)
            elif ext in ['.txt', '.md']:
                return self._extract_from_text(file_bytes)
            elif ext in ['.png', '.jpg', '.jpeg', '.tiff', '.bmp']:
                return self._extract_from_image(file_bytes)
            else:
                raise UnsupportedFileTypeError(f"Unsupported extension: {ext}")
        
        except Exception as e:
            if isinstance(e, (UnsupportedFileTypeError, TextExtractionError)):
                raise
            raise TextExtractionError(f"Failed to extract text from {filename}: {str(e)}")
    
    def _extract_from_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF file."""
        text_parts = []
        
        # Try pdfplumber first (better text extraction)
        try:
            import pdfplumber
            with io.BytesIO(file_bytes) as pdf_file:
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_parts.append(text)
            
            if text_parts:
                return '\n\n'.join(text_parts)
        except ImportError:
            pass
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {e}, trying PyPDF2")
        
        # Fallback to PyPDF2
        try:
            import PyPDF2
            with io.BytesIO(file_bytes) as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
            
            if text_parts:
                return '\n\n'.join(text_parts)
        except Exception as e:
            raise TextExtractionError(f"Failed to extract text from PDF: {e}")
        
        if not text_parts:
            raise TextExtractionError("No text content found in PDF")
        
        return '\n\n'.join(text_parts)
    
    def _extract_from_docx(self, file_bytes: bytes) -> str:
        """Extract text from Word document."""
        try:
            import docx
            with io.BytesIO(file_bytes) as docx_file:
                doc = docx.Document(docx_file)
                text_parts = []
                
                # Extract paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                        if row_text.strip():
                            text_parts.append(row_text)
                
                if not text_parts:
                    raise TextExtractionError("No text content found in Word document")
                
                return '\n\n'.join(text_parts)
        
        except Exception as e:
            raise TextExtractionError(f"Failed to extract text from Word document: {e}")
    
    def _extract_from_text(self, file_bytes: bytes) -> str:
        """Extract text from plain text file."""
        try:
            # Try different encodings
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = file_bytes.decode(encoding)
                    if text.strip():
                        return text
                except UnicodeDecodeError:
                    continue
            
            raise TextExtractionError("Could not decode text file with supported encodings")
        
        except Exception as e:
            raise TextExtractionError(f"Failed to extract text from text file: {e}")
    
    def _extract_from_image(self, file_bytes: bytes) -> str:
        """Extract text from image using OCR."""
        try:
            import pytesseract
            from PIL import Image
            
            with io.BytesIO(file_bytes) as image_file:
                image = Image.open(image_file)
                
                # Perform OCR
                text = pytesseract.image_to_string(image)
                
                if not text.strip():
                    raise TextExtractionError("No text found in image (OCR returned empty)")
                
                return text.strip()
        
        except Exception as e:
            raise TextExtractionError(f"Failed to extract text from image (OCR): {e}")
    
    def parse_and_structure(
        self,
        file_content: Union[bytes, BinaryIO],
        filename: str,
        mime_type: Optional[str] = None,
        use_llm: bool = True
    ) -> Dict[str, Any]:
        """
        Parse a document and return structured content.
        
        Args:
            file_content: File content as bytes or file-like object
            filename: Name of the file
            mime_type: Optional MIME type hint
            use_llm: Whether to use LLM for structuring (default: True)
            
        Returns:
            Dictionary containing:
                - raw_text: Extracted text
                - topics: List of identified topics (if LLM used)
                - overall_summary: Summary of content (if LLM used)
                - metadata: File metadata
                
        Raises:
            DocumentParserError: If parsing fails
        """
        logger.info(f"Parsing document: {filename}")
        
        # Extract text
        try:
            raw_text = self.extract_text(file_content, filename, mime_type)
        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            raise
        
        result = {
            'raw_text': raw_text,
            'metadata': {
                'filename': filename,
                'mime_type': mime_type or mimetypes.guess_type(filename)[0],
                'file_extension': Path(filename).suffix.lower(),
                'text_length': len(raw_text),
                'word_count': len(raw_text.split())
            }
        }
        
        # Structure with LLM if requested
        if use_llm:
            try:
                structured_data = self._structure_with_llm(raw_text)
                result.update(structured_data)
                logger.info(f"Successfully structured content with LLM")
            except Exception as e:
                logger.warning(f"LLM structuring failed: {e}")
                result['llm_error'] = str(e)
                # Return basic structure if LLM fails
                result['topics'] = []
                result['overall_summary'] = raw_text[:500] + '...' if len(raw_text) > 500 else raw_text
        else:
            # Basic structure without LLM
            result['topics'] = []
            result['overall_summary'] = raw_text[:500] + '...' if len(raw_text) > 500 else raw_text
        
        return result
    
    def _structure_with_llm(self, text: str) -> Dict[str, Any]:
        """
        Use LLM to structure extracted text.
        
        Args:
            text: Extracted text content
            
        Returns:
            Structured data with topics, concepts, etc.
        """
        if self.llm_provider is None:
            # Import here to avoid circular dependency
            try:
                from app.services.llm_adapter import get_default_provider
                self.llm_provider = get_default_provider()
            except Exception as e:
                # If no provider or API key configured, gracefully fall back
                logger.warning(f"LLM provider unavailable, using heuristic extraction: {e}")
                return self._extract_basic_structure(text)
        
        # Truncate text if too long (keep first 15000 chars for better LLM processing)
        if len(text) > 15000:
            logger.info(f"Text too long ({len(text)} chars), truncating to 15000 for LLM processing")
            text = text[:15000] + "\n\n[Content truncated for processing...]"
        
        # Use LLM to parse and structure the document
        try:
            structured_data = self.llm_provider.parse_document(text)
            
            # Validate that topics were extracted
            if not structured_data.get('topics'):
                logger.warning("LLM returned no topics, attempting to extract from text structure")
                # Try to extract basic structure from text
                structured_data = self._extract_basic_structure(text)
            
            return structured_data
        except Exception as e:
            logger.error(f"LLM structuring failed: {e}")
            # Fallback to basic extraction
            return self._extract_basic_structure(text)
    
    def _extract_basic_structure(self, text: str) -> Dict[str, Any]:
        """
        Extract basic topic structure from text when LLM fails.
        Uses heuristics to identify topics from headings, numbered sections, etc.
        
        Args:
            text: Raw text content
            
        Returns:
            Basic structured data with topics
        """
        import re
        
        topics = []
        lines = text.split('\n')
        
        # Pattern 1: Look for numbered sections (1., 2., etc.)
        section_pattern = r'^\s*\d+\.\s+(.+)$'
        
        # Pattern 2: Look for UPPERCASE HEADINGS
        uppercase_pattern = r'^([A-Z][A-Z\s]{10,})$'
        
        # Pattern 3: Look for Chapter/Section headers
        chapter_pattern = r'^(Chapter|Section|Topic|Unit|Lesson)\s+\d+:?\s*(.+)$'
        
        current_topic_lines = []
        current_topic_name = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # Check for section headers
            match_section = re.match(section_pattern, line)
            match_uppercase = re.match(uppercase_pattern, line)
            match_chapter = re.match(chapter_pattern, line, re.IGNORECASE)
            
            if match_section or match_uppercase or match_chapter:
                # Save previous topic if exists
                if current_topic_name and current_topic_lines:
                    topics.append({
                        'name': current_topic_name,
                        'description': ' '.join(current_topic_lines[:3]),  # First 3 lines as description
                        'difficulty': 'intermediate',
                        'key_concepts': self._extract_key_concepts(' '.join(current_topic_lines[:10]))
                    })
                    current_topic_lines = []
                
                # Start new topic
                if match_section:
                    current_topic_name = match_section.group(1).strip()
                elif match_uppercase:
                    current_topic_name = match_uppercase.group(1).strip().title()
                elif match_chapter:
                    current_topic_name = match_chapter.group(2).strip()
            else:
                # Accumulate content for current topic
                if current_topic_name:
                    current_topic_lines.append(line)
        
        # Add last topic
        if current_topic_name and current_topic_lines:
            topics.append({
                'name': current_topic_name,
                'description': ' '.join(current_topic_lines[:3]),
                'difficulty': 'intermediate',
                'key_concepts': self._extract_key_concepts(' '.join(current_topic_lines[:10]))
            })
        
        # If no topics found, create a generic one based on text content
        if not topics:
            # Look for key terms that might indicate the subject
            subject_keywords = {
                'probability': ['probability', 'random', 'expected', 'variance', 'distribution'],
                'calculus': ['derivative', 'integral', 'limit', 'continuous', 'differential'],
                'algebra': ['equation', 'polynomial', 'matrix', 'linear', 'quadratic'],
                'statistics': ['mean', 'median', 'standard deviation', 'sample', 'population'],
                'physics': ['force', 'energy', 'velocity', 'momentum', 'mass'],
                'chemistry': ['element', 'compound', 'reaction', 'molecule', 'atom']
            }
            
            text_lower = text.lower()
            detected_subject = None
            max_count = 0
            
            for subject, keywords in subject_keywords.items():
                count = sum(1 for keyword in keywords if keyword in text_lower)
                if count > max_count:
                    max_count = count
                    detected_subject = subject
            
            topic_name = detected_subject.title() if detected_subject else "General Topics"
            
            topics.append({
                'name': topic_name,
                'description': f'Study material covering {topic_name.lower()} concepts',
                'difficulty': 'intermediate',
                'key_concepts': self._extract_key_concepts(text[:1000])
            })
        
        return {
            'topics': topics,
            'overall_summary': text[:500] + '...' if len(text) > 500 else text
        }
    
    def _extract_key_concepts(self, text: str) -> list:
        """
        Extract key concepts from text using simple heuristics.
        
        Args:
            text: Text snippet to extract concepts from
            
        Returns:
            List of key concept strings
        """
        import re
        
        # Look for capitalized terms, formulas, and key phrases
        concepts = set()
        
        # Pattern for capitalized terms (2+ words)
        cap_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)+)\b'
        concepts.update(re.findall(cap_pattern, text))
        
        # Pattern for mathematical formulas/notation
        formula_pattern = r'\b[A-Z]\([A-Z]\)|[A-Z]\s*=\s*'
        if re.search(formula_pattern, text):
            concepts.add('Mathematical Formulas')
        
        # Limit to 10 most meaningful concepts
        concepts = list(concepts)[:10]
        
        # If no concepts found, return generic ones
        if not concepts:
            concepts = ['Key Definitions', 'Important Theorems', 'Practice Problems']
        
        return concepts
    
    def parse_file_path(self, file_path: str, use_llm: bool = True) -> Dict[str, Any]:
        """
        Parse a document from a file path.
        
        Args:
            file_path: Path to the file
            use_llm: Whether to use LLM for structuring
            
        Returns:
            Structured document data
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        with open(file_path, 'rb') as f:
            return self.parse_and_structure(
                file_content=f,
                filename=file_path.name,
                mime_type=mimetypes.guess_type(str(file_path))[0],
                use_llm=use_llm
            )
    
    def get_supported_formats(self) -> Dict[str, Dict[str, Any]]:
        """
        Get information about supported file formats.
        
        Returns:
            Dictionary of supported formats with their status
        """
        formats = {
            'pdf': {
                'extensions': ['.pdf'],
                'mime_types': ['application/pdf'],
                'supported': self.has_pdf_support,
                'description': 'PDF documents'
            },
            'word': {
                'extensions': ['.docx', '.doc'],
                'mime_types': [
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                    'application/msword'
                ],
                'supported': self.has_docx_support,
                'description': 'Microsoft Word documents'
            },
            'text': {
                'extensions': ['.txt', '.md'],
                'mime_types': ['text/plain', 'text/markdown'],
                'supported': True,
                'description': 'Plain text and Markdown files'
            },
            'image': {
                'extensions': ['.png', '.jpg', '.jpeg', '.tiff', '.bmp'],
                'mime_types': ['image/png', 'image/jpeg', 'image/tiff', 'image/bmp'],
                'supported': self.has_ocr_support,
                'description': 'Images with OCR text extraction'
            }
        }
        
        return formats


# ============================================================================
# CONVENIENCE FUNCTIONS
# ============================================================================

def parse_document(
    file_content: Union[bytes, BinaryIO],
    filename: str,
    mime_type: Optional[str] = None,
    use_llm: bool = True,
    llm_provider=None
) -> Dict[str, Any]:
    """
    Convenience function to parse a document.
    
    Args:
        file_content: File content as bytes or file-like object
        filename: Name of the file
        mime_type: Optional MIME type hint
        use_llm: Whether to use LLM for structuring
        llm_provider: Optional LLM provider instance
        
    Returns:
        Structured document data
    """
    parser = DocumentParser(llm_provider=llm_provider)
    return parser.parse_and_structure(file_content, filename, mime_type, use_llm)


def parse_file(file_path: str, use_llm: bool = True, llm_provider=None) -> Dict[str, Any]:
    """
    Convenience function to parse a document from file path.
    
    Args:
        file_path: Path to the file
        use_llm: Whether to use LLM for structuring
        llm_provider: Optional LLM provider instance
        
    Returns:
        Structured document data
    """
    parser = DocumentParser(llm_provider=llm_provider)
    return parser.parse_file_path(file_path, use_llm)


def get_supported_formats() -> Dict[str, Dict[str, Any]]:
    """
    Get information about supported file formats.
    
    Returns:
        Dictionary of supported formats
    """
    parser = DocumentParser()
    return parser.get_supported_formats()
